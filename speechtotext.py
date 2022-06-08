from vosk import Model, SpkModel, KaldiRecognizer
import sys
import wave
import json
import os
import numpy as np
import docx
import subprocess
import PySimpleGUI as sg
# import truecaser
import spacy, re

# UI
layout = [  [sg.Text('Wähle ein MP3 oder .wav File aus:')],
            [sg.In() ,sg.FileBrowse(file_types=(("Audio Files", "*.wav *.mp3, .m4a"),))],
            [sg.Button('Start'), sg.Button('Abbruch')]]

window = sg.Window('Automatische Sprachtranskription', layout)

# truecaser: receives text and uses
def truecase(text):
    nlp = spacy.load('de_core_news_sm')
    doc = nlp(text)
    tagged_sent = [(w.text, w.tag_) for w in doc]
    normalized_sent = [w.capitalize() if t in ["NN","NNS"] else w for (w,t) in tagged_sent]
    normalized_sent[0] = normalized_sent[0].capitalize()
    truecased_text = re.sub(" (?=[\.,'!?:;])", "", ' '.join(normalized_sent))
    return truecased_text

# transcripter: receives audio file in wav, mp3 or m4a format and saves transcript in word file
def transcript_audio (audio):
    model_path = "model"
    spk_model_path = "model-spk"

    # define doc
    doc = docx.Document()
    doc.add_heading('Transcript')

    if not os.path.exists(model_path):
        print(
            "Please download the model from https://alphacephei.com/vosk/models and unpack as {} in the current folder.".format(
                model_path))
        exit(1)

    if not os.path.exists(spk_model_path):
        print(
            "Please download the speaker model from https://alphacephei.com/vosk/models and unpack as {} in the current folder.".format(
                spk_model_path))
        exit(1)

    sample_rate = 16000

    # convert to wave file
    process = subprocess.Popen(['ffmpeg', '-loglevel', 'quiet', '-i',
                                audio,
                                '-ar', str(sample_rate), '-ac', '1', '-f', 's16le', '-'],
                               stdout=subprocess.PIPE)

    model = Model(model_path)
    rec = KaldiRecognizer(model, sample_rate)

    # read data and adds paragraphs to word document
    while True:
        data = process.stdout.read(4000)
        if len(data) == 0:
            break
        if rec.AcceptWaveform(data):
            res = json.loads(rec.Result())
            print(res['text'])
            if res['text'] != "":
                doc.add_paragraph(truecase(res['text']))

    doc.save("C:/Transcripts/test.docx")
    return

# UI window manager
while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED or event == 'Abbruch':
        break
    if event == 'Start':
        transcript_audio(values[0])

window.close()
