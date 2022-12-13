import whisper
import sys
import docx

# Cd Documents\Study\Medizinische Informatik\WS22_23\Masterarbeit\Whisper
# python main.py Nr2_Interview.wav


# define doc
doc = docx.Document()
doc.add_heading('Transcript')

model = whisper.load_model("large")
result = model.transcribe(sys.argv[1], language='german')
print(result["text"])
doc.add_paragraph(result["text"])

#doc.save("C:/Users/jaeger/Transcripts/")
doc.save("C:/Users/jaeger/Transcripts/" + sys.argv[1] + ".docx")
