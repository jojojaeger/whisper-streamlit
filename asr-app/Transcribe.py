import streamlit as st
from transcriber import Transcription
import docx
from datetime import datetime
import pathlib
import io

# App wide config
st.set_page_config(
    page_title="Automatic Transcription",
    layout="wide",
    page_icon="💬"
)
# App sidebar for uplading audio files
with st.sidebar.form("input_form"):
    input_files = st.file_uploader(
        "Files", type=["mp4", "m4a", "mp3", "wav"], accept_multiple_files=True)

    whisper_model = st.selectbox("Whisper model", options=[
        "tiny", "base", "small", "medium", "large"], index=4)

    transcribe = st.form_submit_button(label="Transcribe")

if transcribe:
    if input_files:
        st.session_state.transcription = Transcription(
            input_files)
    else:
        st.error("Please upload a file")

# Transcribe on click
    st.session_state.transcription.transcribe(
        whisper_model
    )

# If there is a transcription, render it. If not, display instructions
if "transcription" in st.session_state:

    doc = docx.Document()

    for output in st.session_state.transcription.output:
        save_dir = str(pathlib.Path(__file__).parent.absolute()
                       ) + "/transcripts/"
        st.markdown(
            f"#### Transcription of {output['name']}")
        st.markdown(
            f"_(whisper model:_`{whisper_model}` -  _language:_ `{output['language']}`)")
        timestamp = -1
        text = ""
        with st.expander("Transcript"):
            for idx, segment in enumerate(output['segments']):
                for s in output['segments'][idx]['whole_word_timestamps']:
                    # Check for pauses in speech longer than 3s
                    #if timestamp != -1 and int(s['timestamp'] - timestamp) >= 3:
                    #    pause = int(s['timestamp'] - timestamp)
                    #    pause_str = "{" + str(pause) + "sek" + "}"
                    #    text += str(f"""{" "}{"."*pause}{pause_str}""")
                    #timestamp = s['timestamp']
                    text += (s['word'])
                    # Insert line break when there is a punctuation mark
                    if "!" in s['word'] or "?" in s['word'] or '.' in s['word']:
                        if not any(i.isdigit() for i in s['word']):
                            any(i.isdigit() for i in s)
                            doc.add_paragraph(text)
                            st.markdown(text)
                            text = ""
            doc.add_paragraph(text)
            st.markdown(text)
        # Save transcript as docx. in local folder
        file_name = output['name'] + "-" + whisper_model + \
            "-" + datetime.today().strftime('%d-%m-%y') + ".docx"
        doc.save(save_dir + file_name)

        bio = io.BytesIO()
        doc.save(bio)
        st.download_button(
            label="Download Transcript",
            data=bio.getvalue(),
            file_name=file_name,
            mime="docx"
        )


else:
    st.header('Introduction')
    st.markdown(
        '- This project was created as part of the master thesis by Johanna Jäger (n51824549@students.meduniwien.ac.at)')
    st.markdown(
        '- The Transcription is done with OpenAI Whisper (https://openai.com/blog/whisper/)')
    st.markdown(
        '- Every file is only saved **locally** - nothing will be uploaded to the cloud')
    st.header('How to use this app')
    st.markdown(
        '1) Select the files you want to transcribe (multiple files possible)')
    st.markdown(
        '2) Choose a model (_large_ for having the most accurate transcription) and click the button')
    st.markdown(
        '3) Check out the resulted transcripts in the _transcripts_-Folder of this directory')
